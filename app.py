import streamlit as st
import os
from google import genai
from PIL import Image
from docx import Document
import io
import socket
import qrcode
import re

API_KEY = "AQ.Ab8RN6KytEbdudB-UQMK_oLPtchUwEWBs9qB7KqzY8ny5a8Xag"
client = genai.Client(api_key=API_KEY)

def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('10.255.255.255', 1))
        IP = s.getsockname()[0]
    except Exception:
        IP = '127.0.0.1'
    finally:
        s.close()
    return IP

def get_sort_key(file):
    name = os.path.splitext(file.name)[0]
    try:
        return int(name) 
    except ValueError:
        return name 

def clean_text(text):
    text = text.replace('**', '')
    text = text.replace('$', '')
    return text.strip()

def analyze_image_to_structured_data(image):
    # Streamlit Cloud 환경에서 400 ClientError가 나지 않도록 바이트로 변환
    buf = io.BytesIO()
    image.save(buf, format="JPEG")
    image_bytes = buf.getvalue()
    
    prompt = """
    이 이미지에 있는 영어 텍스트와 문장을 완벽하게 분석해줘.
    여러 문장이나 문단이 있다면 반드시 '===='로 구분해줘.
    반드시 아래 형식을 지켜서 작성해줘. 다른 형식을 절대 쓰지 마.
    
    [SENTENCE] 여기에 원문 문장을 적어줘
    [SYNTAX] 여기에 구문 분석 및 직독직해를 적어줘
    [VOCAB] 여기에 핵심 단어와 뜻을 적어줘
    [TIP] 여기에 아빠의 꿀팁이나 어법 설명을 적어줘
    ====
    """
    
    response = client.models.generate_content(
        model='gemini-3.5-flash',
        contents=[
            {"mime_type": "image/jpeg", "data": image_bytes},
            prompt
        ],
    )
    return response.text

def parse_analyzed_blocks(analyzed_blocks_list):
    parsed_data = []
    for blocks in analyzed_blocks_list:
        for block in blocks.split('===='):
            if not block.strip():
                continue
            
            sent_text, syntax_text, vocab_text, tip_text = "", "", "", ""
            
            sent_match = re.search(r'\[SENTENCE\](.*?)(?=\[SYNTAX\]|\[VOCAB\]|\[TIP\]|$)', block, re.DOTALL)
            syntax_match = re.search(r'\[SYNTAX\](.*?)(?=\[VOCAB\]|\[TIP\]|\[SENTENCE\]|$)', block, re.DOTALL)
            vocab_match = re.search(r'\[VOCAB\](.*?)(?=\[TIP\]|\[SYNTAX\]|\[SENTENCE\]|$)', block, re.DOTALL)
            tip_match = re.search(r'\[TIP\](.*?)(?=\[SENTENCE\]|\[SYNTAX\]|\[VOCAB\]|$)', block, re.DOTALL)
            
            if sent_match: sent_text = sent_match.group(1).strip()
            if syntax_match: syntax_text = syntax_match.group(1).strip()
            if vocab_match: vocab_text = vocab_match.group(1).strip()
            if tip_match: tip_text = tip_match.group(1).strip()
            
            if not sent_text and not syntax_text:
                sent_text = block.strip()
                syntax_text = "분석 완료"
            
            parsed_data.append({
                "sent": clean_text(sent_text or "원문 문장"),
                "syntax": clean_text(syntax_text or "구문 분석"),
                "vocab": clean_text(vocab_text or "핵심 단어"),
                "tip": clean_text(tip_text or "아빠의 꿀팁")
            })
    return parsed_data

def create_word_document(parsed_data):
    doc = Document()
    doc.add_heading('📖 아빠표 AI 영어 독해 분석 보고서', 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '원문 문장 (Original)'
    hdr_cells[1].text = '구문 분석 및 해설 (Analysis & Tip)'
    
    for item in parsed_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item['sent']
        row_cells[1].text = f"[구문 분석 / 직독직해]\n{item['syntax']}\n\n[핵심 어휘]\n{item['vocab']}\n\n[아빠의 꿀팁]\n{item['tip']}"
                
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

st.set_page_config(layout="wide")

with st.sidebar:
    st.markdown("### 📱 스마트폰으로 촬영하기")
    local_ip = get_local_ip()
    local_url = f"http://{local_ip}:8501"
    
    qr = qrcode.make(local_url)
    buf = io.BytesIO()
    qr.save(buf, format="PNG")
    buf.seek(0)
    
    st.image(buf, caption="스마트폰 카메라로 스캔하세요")

st.title("📖 욘다를 위한 아빠표 영어 독해 분석기")

if 'analysis_data' not in st.session_state:
    st.session_state.analysis_data = None
if 'analysis_file' not in st.session_state:
    st.session_state.analysis_file = None

uploaded_files = st.file_uploader("영어 교재나 지문 사진을 업로드하거나 찍어 올리세요. (여러 장 가능)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_files:
    uploaded_files.sort(key=get_sort_key)
    st.info(f"총 {len(uploaded_files)}장의 이미지가 업로드되었습니다. 분석을 시작합니다!")
    
    if st.button("문장 분석 및 워드 파일 생성"):
        all_analyzed_blocks = []
        
        for idx, file in enumerate(uploaded_files):
            with st.spinner(f"[{idx+1}/{len(uploaded_files)}] '{file.name}' 분석 중..."):
                image = Image.open(file)
                analyzed_result = analyze_image_to_structured_data(image)
                all_analyzed_blocks.append(analyzed_result)
        
        st.session_state.analysis_data = parse_analyzed_blocks(all_analyzed_blocks)
        st.session_state.analysis_file = create_word_document(st.session_state.analysis_data)

if st.session_state.analysis_data:
    st.success("✅ 영어 독해 분석이 완료되었습니다!")
    
    st.download_button(
        label="📥 맞춤형 영어 분석 보고서(.docx) 다운로드",
        data=st.session_state.analysis_file,
        file_name="영어독해분석.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    st.subheader("👀 분석 미리보기")
    for item in st.session_state.analysis_data:
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**[원문]**\n{item['sent']}")
        with col2:
            st.markdown(f"**[구문]** {item['syntax']}")
            st.markdown(f"**[어휘]** {item['vocab']}")
            st.markdown(f"**[꿀팁]** {item['tip']}")
        st.divider()
